from django.shortcuts import render
from custApp.models import Customer
from docx import Document
from django.http import HttpResponse
# Create your views here.

def docx_download(request):
    customer_detail_list = Customer.objects.all()
    document = Document()
    document.add_heading('Customer Details')
    for customer in customer_detail_list:
        document.add_paragraph(' Customer Name: {}'.format(customer.customer_name))
        document.add_paragraph(' Customer Code: {}'.format(customer.customer_code))
        document.add_paragraph(' Customer Address: {}'.format(customer.address))
        document.add_paragraph(' Account Balance: {}'.format(customer.account_bal))
        document.add_paragraph(' Bills Receivable: {}'.format(customer.bill))
        document.add_paragraph(' Advances: {}'.format(customer.advance))
        document.add_paragraph()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)
    return response

def index(request):
    return HttpResponse('<h1>Go to http://127.0.0.1:8000/docx to download customer Information</h1>')
